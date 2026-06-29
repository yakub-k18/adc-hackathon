"""Generate sample_hr.docx for demo purposes."""

from pathlib import Path

from docx import Document


def main() -> None:
    doc = Document()
    doc.add_heading("Employee HR Record (Synthetic)", level=1)
    doc.add_paragraph("Employee Name: Ananya Desai")
    doc.add_paragraph("Employee ID: EMP-45291")
    doc.add_paragraph("Work Email: ananya.desai@examplecorp.com")
    doc.add_paragraph("Mobile Phone: 9876012345")
    doc.add_paragraph("PAN: FGHPJ9876K")
    doc.add_paragraph("Residential Address: 42 Lake View Road, Pune, Maharashtra")
    doc.add_paragraph(
        "Notes: This document contains synthetic HR data for ADC prototype testing only."
    )

    output = Path(__file__).parent / "sample_hr.docx"
    doc.save(output)
    print(f"Created {output}")


if __name__ == "__main__":
    main()
